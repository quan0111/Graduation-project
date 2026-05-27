from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


INPUT = Path(r"D:\DATN\document_work\10_12522084_DaoAnhQuan_BaoCao_DA_DaChinh_Final.docx")
OUTPUT = Path(r"D:\DATN\document_work\10_12522084_DaoAnhQuan_BaoCao_DA_DaChinh_Final_ThamKhao.docx")


def iter_all_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def set_paragraph_text(paragraph, text):
    style = paragraph.style
    clear_paragraph(paragraph)
    paragraph.style = style
    paragraph.add_run(text)


def replace_text(doc, replacements):
    for paragraph in iter_all_paragraphs(doc):
        text = paragraph.text
        new_text = text
        for old, new in replacements.items():
            new_text = new_text.replace(old, new)
        if new_text != text:
            set_paragraph_text(paragraph, new_text)


def add_citation(doc, startswith_or_contains, citation, mode="contains"):
    for paragraph in iter_all_paragraphs(doc):
        text = paragraph.text.strip()
        if not text or citation in text:
            continue
        matched = text.startswith(startswith_or_contains) if mode == "startswith" else startswith_or_contains in text
        if matched:
            set_paragraph_text(paragraph, paragraph.text.rstrip() + " " + citation)


def paragraph_by_text(doc, needle, style=None):
    for paragraph in doc.paragraphs:
        if needle in paragraph.text and (style is None or paragraph.style.name == style):
            return paragraph
    raise ValueError(f"Paragraph containing {needle!r} with style {style!r} not found")


def paragraph_index(paragraphs, paragraph):
    for index, candidate in enumerate(paragraphs):
        if candidate._p is paragraph._p:
            return index
    raise ValueError(f"Paragraph {paragraph.text!r} not found in list")


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def add_paragraph_after(doc, cursor, text="", style=None):
    paragraph = doc.add_paragraph()
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    cursor.addnext(paragraph._p)
    return paragraph._p


def set_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_table_after(doc, cursor, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_header(table.rows[0])
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cells[i].text = value
    cursor.addnext(table._tbl)
    return table._tbl


doc = Document(INPUT)

replace_text(
    doc,
    {
        "Prisma Pythonript": "TypeScript",
        "ReactJS": "React",
        "JWT + JWT": "JWT + refresh token",
        "Postgresql": "PostgreSQL",
        "Quỳ trình": "Quy trình",
        "API được bảo vệ bằng JWT và cơ chế xác thực OAuth": "API được bảo vệ bằng JWT, refresh token và phân quyền RBAC",
        "Ngoài ra, hệ thống hỗ trợ đăng nhập thông qua tài khoản Google nhằm đơn giản hóa quá trình xác thực và nâng cao trải nghiệm người dùng.": "Ngoài ra, hệ thống tách luồng đăng nhập storefront và admin bằng email/mật khẩu, JWT access token và refresh token để hạn chế nhầm quyền giữa người mua, seller và quản trị viên.",
        "frontend được kiểm tra bằng TypeScript build.": "frontend được kiểm tra bằng TypeScript build.",
    },
)

endpoint_updates = {
    "POST /api/v1/auth/google:": (
        "POST /api/v1/auth/admin/login:\n"
        "Đăng nhập vào trang quản trị. API sử dụng email/mật khẩu, cấp access token và refresh token trong phạm vi admin để tách riêng phiên quản trị với phiên storefront."
    ),
    "GET /api/v1/auth/google/callback:": (
        "GET /api/v1/auth/me:\n"
        "Trả về thông tin người dùng hiện tại ở storefront, gồm id, email, họ tên, vai trò, trạng thái tài khoản và thông tin giỏ hàng cơ bản."
    ),
    "POST /api/v1/auth/OAuth/facebook:": (
        "GET /api/v1/auth/admin/me:\n"
        "Kiểm tra phiên đăng nhập admin. Nếu tài khoản không có vai trò ADMIN, API trả về lỗi 403 để bảo vệ khu vực quản trị."
    ),
}
for paragraph in iter_all_paragraphs(doc):
    stripped = paragraph.text.strip()
    for prefix, replacement in endpoint_updates.items():
        if stripped.startswith(prefix):
            set_paragraph_text(paragraph, replacement)

citation_rules = [
    ("FastAPI: framework Python", "[TL01]", "startswith"),
    ("Prisma Python: ORM", "[TL02][TL03]", "startswith"),
    ("PostgreSQL: cơ sở dữ liệu", "[TL03]", "startswith"),
    ("JWT và refresh token", "[TL09][TL10]", "startswith"),
    ("Cloudinary: lưu trữ ảnh", "[TL11]", "startswith"),
    ("AI/ML stack", "[TL12][TL13][TL14][TL15]", "startswith"),
    ("React 19 + TypeScript + Vite", "[TL04][TL05]", "startswith"),
    ("Tailwind CSS: xây dựng giao diện", "[TL16]", "startswith"),
    ("React Router:", "[TL06]", "startswith"),
    ("TanStack Query và Axios", "[TL07]", "startswith"),
    ("Zustand:", "[TL08]", "startswith"),
    ("Recharts:", "[TL17]", "startswith"),
    ("Web API được xây dựng bằng FastAPI", "[TL01][TL02][TL03][NB01]", "contains"),
    ("Xác thực bằng JWT access token", "[TL09][TL10][NB01]", "contains"),
    ("Upload media: hệ thống dùng Cloudinary", "[TL11][NB01]", "contains"),
    ("Tích hợp thành công chức năng gợi ý sản phẩm", "[TL12][TL15][NB01]", "contains"),
    ("Backend được xây dựng bằng FastAPI", "[TL01][TL02][TL03][NB01]", "contains"),
    ("Frontend được xây dựng bằng React 19", "[TL04][TL05][TL06][TL07][TL08][TL16][NB02]", "contains"),
    ("Cơ sở dữ liệu chính là PostgreSQL", "[TL03][NB01]", "contains"),
]
for needle, citation, mode in citation_rules:
    add_citation(doc, needle, citation, mode)

style_fixes = [
    ("2.2 Công nghệ áp dụng", "MucLon"),
    ("Prisma Python: ORM", "Nội dung"),
    ("Tailwind CSS: xây dựng giao diện", "Nội dung"),
    ("2.2.3 Công nghệ hỗ trợ khác", "MucNho"),
    ("Git/GitHub: quản lý phiên bản", "Nội dung"),
    ("Swagger/OpenAPI: kiểm thử", "Nội dung"),
    ("Tích hợp thanh toán và hoàn tiền ở chế độ production", "Nội dung"),
]
for paragraph in iter_all_paragraphs(doc):
    stripped = paragraph.text.strip()
    for prefix, style_name in style_fixes:
        if stripped.startswith(prefix):
            paragraph.style = style_name

# Add a short theoretical/source section before chapter 3.
chapter3 = paragraph_by_text(doc, "PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG")
previous = chapter3._p.getprevious()
cursor = previous if previous is not None else chapter3._p
for text, style in [
    ("2.3 Cơ sở lựa chọn công nghệ và nguồn tham khảo", "MucNho"),
    (
        "Các nội dung kỹ thuật trong báo cáo không chỉ mô tả theo kinh nghiệm triển khai mà còn được đối chiếu với tài liệu chính thức của framework/thư viện và mã nguồn hiện tại của đồ án. Cách trình bày này giúp người đọc phân biệt phần cơ sở lý thuyết, phần quyết định thiết kế và phần đã được triển khai thực tế.",
        "Nội dung",
    ),
    (
        "Đối với backend, FastAPI phù hợp để xây dựng REST API có tài liệu OpenAPI/Swagger, Prisma Python giúp ánh xạ schema với PostgreSQL, JWT/refresh token kết hợp RBAC giúp kiểm soát quyền truy cập theo vai trò, còn Cloudinary hỗ trợ lưu trữ media sản phẩm và review [TL01][TL02][TL03][TL09][TL10][TL11].",
        "Nội dung",
    ),
    (
        "Đối với frontend, React được dùng để xây dựng giao diện theo component, Vite hỗ trợ môi trường phát triển nhanh, React Router quản lý điều hướng, TanStack Query xử lý dữ liệu server/cache, Zustand xử lý một số trạng thái client và Tailwind CSS hỗ trợ thiết kế responsive [TL04][TL05][TL06][TL07][TL08][TL16].",
        "Nội dung",
    ),
    (
        "Đối với recommendation, phần lý thuyết dựa trên nhóm tài liệu về hệ gợi ý, đánh giá ranking và Learning-to-Rank. Vì dữ liệu hành vi trong đồ án còn nhỏ, báo cáo cần nêu rõ rằng HitRate@K, NDCG@K, CTR và Conversion chỉ có ý nghĩa khi tập dữ liệu đủ lớn và có quy trình holdout/A-B testing phù hợp [TL12][TL13][TL14][TL15].",
        "Nội dung",
    ),
    (
        "Các nghiệp vụ đặc thù của đồ án như đơn hàng tách theo shop, trạng thái vận chuyển riêng từng package, hoa hồng nền tảng, payout seller, review sau mua và dữ liệu hành vi người dùng được đối chiếu trực tiếp với schema.prisma, các service backend và các màn hình user-fe/admin-FE hiện có [NB01][NB02].",
        "Nội dung",
    ),
]:
    cursor = add_paragraph_after(doc, cursor, text, style)

# Add a finance design subsection before the architecture subsection.
design_heading = paragraph_by_text(doc, "Thiết kế hệ thống", style="Mục lớn")
cursor = design_heading._p
for text, style in [
    ("Thiết kế tài chính, hoa hồng và payout", "Tiểu mục"),
    (
        "Trong hệ thống thương mại điện tử đa người bán, doanh thu của đơn hàng cần được tách theo từng shop thay vì chỉ tính một tổng đơn chung. Mỗi shop có thể có trạng thái giao hàng, hoàn tiền và đối soát khác nhau, do đó hệ thống sử dụng OrderShopPackage để quản lý package riêng cho từng seller [NB01].",
        "Nội dung",
    ),
    (
        "Khi một package hoàn tất, hệ thống tính phí nền tảng theo cấu hình hoa hồng từ 3% đến 7% tùy đơn hàng/danh mục/cấu hình shop. Khoản phí này được lưu vào PlatformCommission, sau đó phần còn lại được ghi nhận là doanh thu có thể đối soát cho seller [NB01].",
        "Nội dung",
    ),
    (
        "Luồng rút tiền của seller sử dụng SellerPayout. Khi seller gửi yêu cầu, số tiền được đưa vào trạng thái chờ xử lý; admin có quyền duyệt hoặc từ chối. Sau khi duyệt, hệ thống cập nhật trạng thái payout và ghi nhận biến động tài chính để đối soát về sau [NB01].",
        "Nội dung",
    ),
    (
        "Thiết kế này giúp admin theo dõi được tổng phí nền tảng, seller biết rõ doanh thu thực nhận, đồng thời hạn chế việc rút vượt số dư khả dụng. Đây là phần cần nhấn mạnh trong báo cáo vì nó thể hiện đặc trưng marketplace, khác với website bán hàng một chủ thông thường [NB01][NB02].",
        "Nội dung",
    ),
]:
    cursor = add_paragraph_after(doc, cursor, text, style)

# Replace the existing reference placeholder with a real reference map and bibliography.
ref_heading = paragraph_by_text(doc, "TÀI LIỆU THAM KHẢO")
appendix_heading = paragraph_by_text(doc, "PHỤ LỤC")
paragraphs = list(doc.paragraphs)
ref_idx = paragraph_index(paragraphs, ref_heading)
appendix_idx = paragraph_index(paragraphs, appendix_heading)
for paragraph in paragraphs[ref_idx + 1 : appendix_idx]:
    delete_paragraph(paragraph)

cursor = ref_heading._p
reference_intro = (
    "Ghi chú: [TLxx] là tài liệu bên ngoài dùng làm cơ sở lý thuyết/công nghệ; "
    "[NBxx] là nguồn nội bộ được đối chiếu từ mã nguồn đồ án hiện tại."
)
cursor = add_paragraph_after(doc, cursor, reference_intro, "Nội dung")
cursor = add_paragraph_after(doc, cursor, "Bảng đối chiếu nội dung tham khảo trong báo cáo", "Caption")
cursor = add_table_after(
    doc,
    cursor,
    ["Nội dung trong báo cáo", "Vị trí sử dụng", "Nguồn đối chiếu"],
    [
        (
            "Công nghệ backend: FastAPI, Prisma Python, PostgreSQL, REST API, OpenAPI/Swagger",
            "Mục 2.2.1, 2.3, 4.1",
            "[TL01][TL02][TL03][NB01]",
        ),
        (
            "Xác thực, refresh token, phân quyền Customer/Seller/Admin và xử lý lỗi 403",
            "Mục 2.2.1, 3.2, 4.1.2",
            "[TL09][TL10][NB01]",
        ),
        (
            "Frontend: React, TypeScript, Vite, React Router, TanStack Query, Zustand, Tailwind CSS",
            "Mục 2.2.2, 3.3, 4.3",
            "[TL04][TL05][TL06][TL07][TL08][TL16][NB02]",
        ),
        (
            "Upload ảnh sản phẩm, ảnh variant, media review và bằng chứng trả hàng",
            "Mục 2.2.1, 4.2",
            "[TL11][NB01][NB02]",
        ),
        (
            "Gợi ý sản phẩm bằng AI, Learning-to-Rank, LightGBM/XGBoost, HitRate@K, NDCG@K",
            "Mục 1.4, 2.3, 4.2, 5.1",
            "[TL12][TL13][TL14][TL15][NB01]",
        ),
        (
            "Đơn hàng nhiều shop, tracking theo package, hoa hồng nền tảng và payout seller",
            "Mục 3.1, 3.2, 3.4, 4.2",
            "[NB01][NB02]",
        ),
        (
            "Dashboard admin/seller, biểu đồ doanh thu, analytics và báo cáo vận hành",
            "Mục 2.2.2, 3.3, 4.2",
            "[TL17][NB02]",
        ),
    ],
)
cursor = add_paragraph_after(doc, cursor, "Danh mục tài liệu tham khảo", "MucNho")

references = [
    "[TL01] FastAPI. Features - FastAPI. https://fastapi.tiangolo.com/features/",
    "[TL02] Robert Craigie. Prisma Client Python Documentation. https://prisma-client-py.readthedocs.io/en/stable/",
    "[TL03] PostgreSQL Global Development Group. PostgreSQL Documentation: What Is PostgreSQL? và Transactions. https://www.postgresql.org/docs/current/",
    "[TL04] React. Describing the UI - React Documentation. https://react.dev/learn/describing-the-ui",
    "[TL05] Vite. Getting Started - Vite Guide. https://vite.dev/guide/",
    "[TL06] React Router. Routing - React Router Documentation. https://reactrouter.com/start/declarative/routing",
    "[TL07] TanStack Query. Overview - TanStack Query React Docs. https://tanstack.com/query/latest/docs/framework/react/overview",
    "[TL08] Zustand. Introduction - Zustand Documentation. https://zustand.docs.pmnd.rs/getting-started/introduction",
    "[TL09] Jones, M., Bradley, J., Sakimura, N. RFC 7519: JSON Web Token (JWT). IETF, 2015. https://datatracker.ietf.org/doc/rfc7519/",
    "[TL10] OWASP Cheat Sheet Series. Authorization Cheat Sheet. https://cheatsheetseries.owasp.org/cheatsheets/Authorization_Cheat_Sheet.html",
    "[TL11] Cloudinary. Programmatically Uploading Images, Videos, and Other Files. https://cloudinary.com/documentation/upload_images",
    "[TL12] Ricci, F., Rokach, L., Shapira, B. Recommender Systems Handbook, 2nd edition. Springer, 2015. https://link.springer.com/book/10.1007/978-1-4899-7637-6",
    "[TL13] LightGBM. Parameters and LGBMRanker Documentation. https://lightgbm.readthedocs.io/en/latest/Parameters.html",
    "[TL14] XGBoost. Learning to Rank Tutorial. https://xgboost.readthedocs.io/en/stable/tutorials/learning_to_rank.html",
    "[TL15] Herlocker, J. L., Konstan, J. A., Terveen, L. G., Riedl, J. T. Evaluating Recommender Systems. ACM TOIS, 2004. https://www.microsoft.com/en-us/research/publication/evaluating-recommender-systems/",
    "[TL16] Tailwind CSS. Styling with utility classes. https://tailwindcss.com/docs/styling-with-utility-classes",
    "[TL17] Recharts. Recharts Documentation. https://recharts.github.io/en-US/",
    "[NB01] Mã nguồn backend đồ án: D:\\DATN\\BE\\src\\modules và D:\\DATN\\BE\\prisma\\schema.prisma, đối chiếu ngày 27/05/2026.",
    "[NB02] Mã nguồn frontend đồ án: D:\\DATN\\user-fe và D:\\DATN\\admin-FE, đối chiếu ngày 27/05/2026.",
]
for ref in references:
    cursor = add_paragraph_after(doc, cursor, ref, "Bibliography")

# Replace appendix placeholders with useful project appendices.
appendix_heading = paragraph_by_text(doc, "PHỤ LỤC")
all_paragraphs = list(doc.paragraphs)
appendix_idx = paragraph_index(all_paragraphs, appendix_heading)
for paragraph in all_paragraphs[appendix_idx + 1 :]:
    delete_paragraph(paragraph)

cursor = appendix_heading._p
appendix_items = [
    ("Phụ lục A. Cấu trúc mã nguồn chính", "Phụ mục"),
    ("Backend: D:\\DATN\\BE gồm FastAPI app, prisma schema, router/service/schema theo từng module nghiệp vụ như auth, product, order, shipment, return_request, finance, analytics, seller, admin, upload và support.", "Nội dung"),
    ("Frontend người dùng/seller: D:\\DATN\\user-fe gồm các trang mua hàng, catalog, product detail, cart, checkout, order tracking, wishlist, review, seller center và finance.", "Nội dung"),
    ("Frontend quản trị: D:\\DATN\\admin-FE gồm dashboard, quản lý user/shop/product/order, moderation, analytics/recommendation, finance commission và duyệt payout.", "Nội dung"),
    ("Phụ lục B. Các luồng nghiệp vụ cần minh họa khi bảo vệ", "Phụ mục"),
    ("Luồng checkout nhiều shop: người dùng đặt một đơn tổng, hệ thống tách OrderShopPackage theo shop, mỗi package có trạng thái vận chuyển và xác nhận nhận hàng riêng.", "Nội dung"),
    ("Luồng tài chính: khi package hoàn tất, hệ thống tính PlatformCommission theo tỉ lệ 3-7%, cập nhật số dư seller và cho phép seller tạo SellerPayout để admin duyệt.", "Nội dung"),
    ("Luồng recommendation: hệ thống ghi nhận UserBehavior, train/evaluate mô hình, tính HitRate@K, NDCG@K, CTR, Conversion và trả danh sách sản phẩm gợi ý cho frontend.", "Nội dung"),
]
for text, style in appendix_items:
    cursor = add_paragraph_after(doc, cursor, text, style)

doc.save(OUTPUT)
print(OUTPUT)
