from pathlib import Path

from docx import Document


SOURCE = Path(r"D:\DATN\document_work\bao_cao_goc.docx")
OUTPUT = Path(r"D:\DATN\document_work\10_12522084_DaoAnhQuan_BaoCao_DA_DaChinh.docx")


def set_text(paragraph, text: str):
    paragraph.clear()
    if text:
        paragraph.add_run(text)


def delete_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def replace_range(doc, start: int, end: int, texts: list[str]):
    paragraphs = list(doc.paragraphs)
    targets = paragraphs[start : end + 1]
    if not targets:
        return

    for index, text in enumerate(texts):
        if index < len(targets):
            set_text(targets[index], text)
        else:
            break

    for paragraph in reversed(targets[len(texts) :]):
        delete_paragraph(paragraph)


def replace_in_all_text(doc, replacements: dict[str, str]):
    def update_paragraph(paragraph):
        text = paragraph.text
        updated = text
        for old, new in replacements.items():
            updated = updated.replace(old, new)
        if updated != text:
            set_text(paragraph, updated)

    for paragraph in doc.paragraphs:
        update_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    update_paragraph(paragraph)


def update_table_text(doc):
    # Table 2: glossary table.
    if len(doc.tables) > 2:
        glossary = doc.tables[2]
        rows = [
            ("Từ viết tắt", "Từ đầy đủ", "Giải thích"),
            ("API", "Application Programming Interface", "Giao diện lập trình để frontend giao tiếp với backend."),
            ("JWT", "JSON Web Token", "Token dùng để xác thực và phân quyền người dùng."),
            ("RBAC", "Role-Based Access Control", "Cơ chế phân quyền theo vai trò Customer, Seller và Admin."),
            ("ORM", "Object-Relational Mapping", "Công cụ ánh xạ dữ liệu giữa mã nguồn Python và PostgreSQL."),
            ("SKU", "Stock Keeping Unit", "Mã định danh biến thể sản phẩm phục vụ quản lý kho."),
            ("COD", "Cash on Delivery", "Phương thức thanh toán khi nhận hàng."),
            ("VNPay/MoMo", "Cổng thanh toán trực tuyến", "Các phương thức thanh toán online được tích hợp trong hệ thống."),
            ("LTR", "Learning to Rank", "Mô hình học xếp hạng dùng trong hệ thống gợi ý sản phẩm."),
            ("Payout", "Yêu cầu rút tiền", "Quy trình seller gửi yêu cầu rút số dư và admin duyệt thanh toán."),
        ]
        for row_index, row_values in enumerate(rows):
            if row_index >= len(glossary.rows):
                break
            for col_index, value in enumerate(row_values):
                if col_index < len(glossary.rows[row_index].cells):
                    glossary.rows[row_index].cells[col_index].text = value

    # Table 3: scope examples.
    if len(doc.tables) > 3:
        table = doc.tables[3]
        data = [
            ("Tiêu chí", "Phạm vi áp dụng", "Nội dung chính", "Ghi chú"),
            ("Không gian", "Người dùng tại Việt Nam", "Marketplace đa người bán, giao diện tiếng Việt", "Triển khai và kiểm thử trên môi trường cục bộ."),
            ("Thời gian", "03/2026 - 06/2026", "Phân tích, thiết kế, triển khai, kiểm thử và hoàn thiện báo cáo", "Dữ liệu thử nghiệm được seed trong PostgreSQL."),
            ("Khoa học", "Full-stack web và AI", "React + Vite, FastAPI, Prisma, PostgreSQL, Recommendation Engine", "Có tracking hành vi và đánh giá offline."),
            ("Thực tiễn", "Sinh viên, shop nhỏ", "Mua sắm, bán hàng, vận chuyển, hoàn tiền, tài chính seller", "Hỗ trợ seller center và admin center."),
        ]
        for row_index, row_values in enumerate(data):
            if row_index >= len(table.rows):
                break
            for col_index, value in enumerate(row_values):
                if col_index < len(table.rows[row_index].cells):
                    table.rows[row_index].cells[col_index].text = value

    # Table 4: feature overview.
    if len(doc.tables) > 4:
        table = doc.tables[4]
        rows = [
            ("STT", "Chức năng", "Mô tả"),
            ("1", "Đăng ký / đăng nhập", "Người dùng đăng ký, đăng nhập bằng email và mật khẩu; hệ thống cấp JWT access token và refresh token."),
            ("2", "Xem và tìm kiếm sản phẩm", "Hiển thị danh sách sản phẩm, chi tiết sản phẩm, ảnh, biến thể, thuộc tính, đánh giá và shop bán."),
            ("3", "Giỏ hàng và đặt hàng", "Thêm sản phẩm vào giỏ, cập nhật số lượng, áp dụng coupon và tạo đơn hàng."),
            ("4", "Thanh toán", "Hỗ trợ COD, VNPay và MoMo; cập nhật trạng thái thanh toán và đơn hàng."),
            ("5", "Theo dõi đơn đa shop", "Mỗi shop trong cùng một hóa đơn có package và trạng thái vận chuyển riêng."),
            ("6", "Đánh giá và wishlist", "Người dùng lưu sản phẩm yêu thích, đánh giá sản phẩm đã mua và xem review có media."),
            ("7", "Trả hàng / hoàn tiền", "Người mua gửi yêu cầu trả hàng, seller/admin xử lý và cập nhật trạng thái hoàn tiền."),
            ("8", "Seller Center", "Seller quản lý shop, sản phẩm, variant, tồn kho, đơn hàng, vận chuyển, review, coupon, marketing và tài chính."),
            ("9", "Admin Center", "Admin quản lý user, shop, đơn hàng, sản phẩm, kiểm duyệt, return, marketing, analytics, hoa hồng và payout."),
            ("10", "AI Recommendation", "Ghi nhận VIEW/CLICK/ADD_TO_CART/PURCHASE, train model và đánh giá HitRate@K, NDCG@K, CTR, Conversion."),
            ("11", "Tài chính nền tảng", "Tính hoa hồng 3-7% theo đơn, ghi ledger commission, khóa số dư pending payout và duyệt rút tiền."),
            ("12", "Upload media", "Upload ảnh sản phẩm, ảnh variant, bằng chứng return và media review qua Cloudinary."),
            ("13", "Bảo mật và audit", "RBAC theo vai trò, audit log, security incident, refresh token và soft delete dữ liệu quan trọng."),
        ]
        for row_index, row_values in enumerate(rows):
            if row_index >= len(table.rows):
                break
            for col_index, value in enumerate(row_values):
                if col_index < len(table.rows[row_index].cells):
                    table.rows[row_index].cells[col_index].text = value


def main():
    doc = Document(SOURCE)

    # Work bottom-up so paragraph deletions do not invalidate later ranges.
    replace_range(
        doc,
        803,
        831,
        [
            "Hạn chế của đề tài",
            "Mặc dù hệ thống đã hoàn thiện nhiều chức năng nghiệp vụ, đồ án vẫn còn một số giới hạn cần tiếp tục cải thiện:",
            "Dữ liệu hành vi dùng cho mô hình gợi ý vẫn chủ yếu phục vụ demo và kiểm thử; để đánh giá chính xác trong môi trường production cần thu thập dữ liệu thật từ số lượng người dùng lớn hơn.",
            "Luồng payout hiện dừng ở mức duyệt nghiệp vụ trong hệ thống, chưa tích hợp API chuyển khoản ngân hàng thực tế.",
            "Thông báo realtime và chat trực tiếp giữa người mua - người bán có thể tiếp tục được hoàn thiện bằng WebSocket hoặc service message queue.",
            "Hệ thống mới kiểm thử trên môi trường cục bộ, chưa đánh giá tải lớn với số lượng sản phẩm, đơn hàng và truy cập đồng thời cao.",
            "Giao diện hiện tập trung cho tiếng Việt, chưa triển khai đa ngôn ngữ và chưa tối ưu sâu cho SEO.",
            "Hướng phát triển của đề tài",
            "Trong tương lai, hệ thống có thể phát triển theo các hướng sau:",
            "Triển khai production hoàn chỉnh với domain, HTTPS, CI/CD, backup database và giám sát log.",
            "Tích hợp thanh toán và hoàn tiền ở chế độ production, đồng thời kết nối API ngân hàng để tự động hóa payout cho seller.",
            "Bổ sung realtime notification/chat, giúp người mua, seller và admin theo dõi biến động đơn hàng nhanh hơn.",
            "Mở rộng hệ thống gợi ý bằng dữ liệu người dùng thật, thử nghiệm thêm mô hình hybrid recommendation, vector search và A/B testing.",
            "Tối ưu hiệu năng backend, cache Redis, phân trang dữ liệu lớn và tối ưu truy vấn Prisma/PostgreSQL.",
            "Bổ sung đa ngôn ngữ, dark mode và cải thiện trải nghiệm mobile-first.",
            "Mở rộng báo cáo quản trị như phân tích cohort, hiệu quả coupon, hiệu quả banner/flash sale và chất lượng vận hành seller.",
            "Với các hướng phát triển trên, hệ thống có tiềm năng trở thành một nền tảng thương mại điện tử thông minh, có khả năng hỗ trợ đầy đủ vòng đời mua bán, vận hành seller và cá nhân hóa trải nghiệm người dùng.",
        ],
    )

    replace_range(
        doc,
        778,
        783,
        [
            "Triển khai ứng dụng",
            "Hệ thống được thiết kế theo mô hình client-server tách biệt rõ ràng giữa frontend, backend và database. Mã nguồn được chia thành ba phần chính: user-fe cho giao diện người dùng/seller, admin-FE cho quản trị viên và BE cho API backend.",
            "Backend được xây dựng bằng FastAPI, tổ chức theo module nghiệp vụ như auth, product, order, shipment, return_request, finance, analytics, seller, marketing, inventory, support và upload. Backend sử dụng Prisma Python để thao tác với PostgreSQL.",
            "Frontend được xây dựng bằng React 19, TypeScript và Vite. Giao diện sử dụng Tailwind CSS, React Router, TanStack Query, Zustand, Recharts, Sonner và lucide-react để xây dựng trải nghiệm responsive cho user, seller và admin.",
            "Cơ sở dữ liệu chính là PostgreSQL, lưu trữ đầy đủ user, shop, product, variant, order, package theo shop, payment, return, review, wishlist, behavior, commission, payout và audit/security log.",
            "Trong giai đoạn đồ án, hệ thống được vận hành và kiểm thử trên môi trường cục bộ; cấu trúc Docker, nginx và biến môi trường đã được chuẩn bị để có thể triển khai lên môi trường production.",
        ],
    )

    replace_range(
        doc,
        719,
        727,
        [
            "Các chức năng hệ thống",
            "Xác thực và phân quyền: hệ thống hỗ trợ đăng ký, đăng nhập bằng email/mật khẩu, cấp JWT access token và refresh token; frontend user và admin sử dụng luồng đăng nhập riêng để hạn chế nhầm quyền.",
            "Phân quyền theo vai trò: Customer, Seller và Admin. Customer mua hàng và đánh giá; Seller quản lý shop, sản phẩm, đơn, vận chuyển và tài chính; Admin quản trị toàn hệ thống.",
            "Quản lý phiên làm việc: refresh token được lưu và có thể thu hồi khi logout, token hết hạn hoặc tài khoản bị khóa.",
            "Audit và bảo mật: các thao tác quan trọng như đăng nhập, cập nhật đơn hàng, xử lý return, duyệt payout và kiểm duyệt sản phẩm được ghi log để phục vụ truy vết.",
            "Upload media: hệ thống dùng Cloudinary để lưu ảnh sản phẩm, ảnh variant, bằng chứng trả hàng và media review.",
            "Thông báo và hỗ trợ: hệ thống có module notification, support ticket và chatbot hỗ trợ tra cứu thông tin sản phẩm/đơn hàng.",
        ],
    )

    replace_range(
        doc,
        651,
        658,
        [
            "Trong đồ án này, Web API được xây dựng bằng FastAPI (Python) kết hợp Prisma Python để thao tác với cơ sở dữ liệu PostgreSQL. Kiến trúc API tuân theo phong cách RESTful, phân tách rõ router, schema và service cho từng module nghiệp vụ.",
            "Phương pháp xây dựng API",
            "Sử dụng FastAPI làm core framework, mỗi nghiệp vụ được tổ chức thành module riêng gồm router, schema và service.",
            "Áp dụng Prisma Python để mapping model trong schema.prisma với bảng dữ liệu PostgreSQL.",
            "Sử dụng Pydantic schema để validate dữ liệu đầu vào và chuẩn hóa dữ liệu trả về.",
            "Xác thực bằng JWT access token, refresh token và phân quyền RBAC theo Customer, Seller, Admin.",
            "Sử dụng HTTPException, audit log và service layer để xử lý lỗi và nghiệp vụ thống nhất.",
            "Tách các nghiệp vụ phức tạp như checkout, shipment, return, commission, payout và recommendation thành service riêng để dễ kiểm thử và mở rộng.",
        ],
    )

    replace_range(
        doc,
        433,
        473,
        [
            "Các yêu cầu phi chức năng",
            "a) Yêu cầu về hiệu năng",
            "Tốc độ phản hồi: danh sách sản phẩm, chi tiết sản phẩm, giỏ hàng và đơn hàng cần phản hồi nhanh để không làm gián đoạn trải nghiệm mua sắm.",
            "Giải pháp: frontend sử dụng Vite, React Query để cache dữ liệu API; backend tổ chức service theo module, hỗ trợ phân trang, soft delete và cache Redis cho các dữ liệu phù hợp.",
            "Khả năng mở rộng: hệ thống phải có khả năng mở rộng số lượng shop, sản phẩm, biến thể, đơn hàng và hành vi người dùng.",
            "Giải pháp: cơ sở dữ liệu PostgreSQL được thiết kế quan hệ rõ ràng; các bảng quan trọng có index theo userId, shopId, productId, orderId, status và createdAt.",
            "b) Yêu cầu về giao diện người dùng",
            "Giao diện cần thân thiện, dễ thao tác và phù hợp với ba nhóm người dùng: khách mua hàng, seller và admin.",
            "Giải pháp: user-fe và admin-FE được tách riêng; Tailwind CSS, component UI, icon và chart được sử dụng nhất quán để tạo giao diện responsive.",
            "Đối với seller, giao diện cần ưu tiên hiệu quả vận hành như quản lý đơn, vận chuyển, tồn kho, review và tài chính. Đối với admin, giao diện cần hỗ trợ rà soát nhanh các trạng thái vận hành.",
            "c) Yêu cầu về an toàn dữ liệu",
            "Hệ thống phải bảo vệ tài khoản, phân quyền đúng vai trò và tránh cho người dùng truy cập dữ liệu không thuộc quyền của mình.",
            "Giải pháp: backend kiểm tra role ở từng endpoint, seller chỉ thao tác với shop của mình, customer chỉ xác nhận package thuộc đơn của mình và admin có quyền quản trị tập trung.",
            "Các dữ liệu quan trọng như sản phẩm, người dùng, đơn hàng sử dụng soft delete hoặc trạng thái nghiệp vụ thay vì xóa cứng.",
            "d) Yêu cầu về tính đúng đắn nghiệp vụ",
            "Đơn hàng nhiều shop phải được theo dõi theo từng gói hàng của từng shop; mỗi package có trạng thái, mã vận đơn, thời điểm giao và xác nhận nhận hàng riêng.",
            "Phí nền tảng phải được tính minh bạch theo từng order item/package, lưu vào PlatformCommission và tổng hợp trong trang tài chính admin.",
            "Yêu cầu payout phải khóa số dư pending, admin duyệt/từ chối và seller theo dõi trạng thái rút tiền.",
            "e) Yêu cầu về khả năng kiểm thử",
            "Các luồng quan trọng như checkout, trạng thái đơn hàng, recommendation metrics và finance commission cần có test để hạn chế lỗi hồi quy.",
            "Giải pháp: backend sử dụng unittest cho checkout lifecycle, analytics/recommendation và finance service; frontend được kiểm tra bằng TypeScript build.",
        ],
    )

    replace_range(
        doc,
        307,
        329,
        [
            "Đặc tả yêu cầu phần mềm",
            "Hệ thống hướng đến xây dựng một nền tảng thương mại điện tử đa người bán, trong đó khách hàng có thể tìm kiếm, mua hàng và nhận gợi ý sản phẩm; seller có thể vận hành shop; admin có thể quản lý toàn bộ nền tảng.",
            "Các vấn đề nghiệp vụ cần giải quyết bao gồm:",
            "Số lượng sản phẩm và shop lớn khiến người dùng khó tìm sản phẩm phù hợp nếu chỉ duyệt thủ công.",
            "Một đơn hàng có thể chứa sản phẩm từ nhiều shop, do đó trạng thái xử lý và vận chuyển cần được tách theo từng shop package.",
            "Seller cần một trung tâm vận hành riêng để quản lý sản phẩm, variant, tồn kho, đơn hàng, vận chuyển, review, coupon, marketing và tài chính.",
            "Admin cần công cụ kiểm duyệt seller/sản phẩm, xử lý đơn/return, theo dõi hành vi, train/evaluate recommendation, tính hoa hồng và duyệt payout.",
            "Hệ thống cần ghi nhận hành vi người dùng như xem sản phẩm, click, thêm vào giỏ và mua hàng để phục vụ recommendation.",
            "Nghiệp vụ chính của hệ thống bao gồm:",
            "Quản lý tài khoản, địa chỉ, phân quyền Customer/Seller/Admin và bảo mật phiên đăng nhập.",
            "Quản lý catalog sản phẩm gồm category, product, variant, image, variant image, attribute, tag, review và wishlist.",
            "Quản lý giỏ hàng, coupon, flash sale, checkout, thanh toán COD/VNPay/MoMo và theo dõi trạng thái đơn hàng.",
            "Quản lý vận chuyển theo package riêng của từng shop trong cùng một hóa đơn.",
            "Quản lý trả hàng/hoàn tiền với bằng chứng media và luồng duyệt giữa customer, seller, admin.",
            "Quản lý tài chính seller gồm doanh thu hoàn tất, phí nền tảng 3-7%, ledger commission, yêu cầu payout và duyệt payout.",
            "Tích hợp AI Recommendation để cá nhân hóa danh sách gợi ý, đồng thời có dashboard admin để train và đánh giá mô hình.",
        ],
    )

    replace_range(
        doc,
        278,
        303,
        [
            "2.2 Công nghệ áp dụng",
            "Đồ án sử dụng các công nghệ hiện đại, phù hợp với mô hình web full-stack tách biệt frontend, backend, database và AI module.",
            "2.2.1 Công nghệ backend",
            "FastAPI: framework Python dùng để xây dựng REST API có hiệu năng tốt, hỗ trợ OpenAPI/Swagger, dependency injection và dễ tổ chức theo module.",
            "Prisma Python: ORM dùng để ánh xạ schema.prisma với cơ sở dữ liệu PostgreSQL, giúp thao tác dữ liệu rõ ràng và giảm lỗi truy vấn thủ công.",
            "PostgreSQL: cơ sở dữ liệu quan hệ chính, lưu trữ user, shop, product, variant, order, payment, shipment, return, review, wishlist, behavior, commission và payout.",
            "JWT và refresh token: dùng cho xác thực, phân tách phiên storefront/admin và kiểm soát quyền theo vai trò Customer, Seller, Admin.",
            "Cloudinary: lưu trữ ảnh sản phẩm, ảnh variant, media review và bằng chứng trả hàng/hoàn tiền.",
            "AI/ML stack: NumPy, joblib, LightGBM, XGBoost và các model recommendation nội bộ để train, lưu và phục vụ gợi ý sản phẩm.",
            "Redis và APScheduler: hỗ trợ cache, rate limit và tác vụ định kỳ khi triển khai mở rộng.",
            "2.2.2 Công nghệ frontend",
            "React 19 + TypeScript + Vite: xây dựng giao diện theo component, tốc độ phát triển nhanh và tách riêng user-fe với admin-FE.",
            "Tailwind CSS: xây dựng giao diện responsive, nhất quán cho trang người dùng, Seller Center và Admin Center.",
            "React Router: quản lý điều hướng giữa trang chủ, catalog, product detail, cart, order, seller và admin.",
            "TanStack Query và Axios: quản lý trạng thái dữ liệu server, gọi API, cache, refetch và xử lý loading/error.",
            "Zustand: quản lý một số trạng thái toàn cục phía client như phiên người dùng hoặc trạng thái giao diện.",
            "Recharts: hiển thị biểu đồ doanh thu, analytics, finance và các thống kê vận hành.",
            "Sonner và lucide-react: hiển thị thông báo, icon và phản hồi thao tác trên giao diện.",
            "2.2.3 Công nghệ hỗ trợ khác",
            "Git/GitHub: quản lý phiên bản mã nguồn.",
            "Swagger/OpenAPI: kiểm thử và tài liệu hóa API backend.",
            "Docker, nginx và biến môi trường: hỗ trợ đóng gói, reverse proxy và chuẩn bị triển khai production.",
            "Visual Studio Code: môi trường phát triển chính.",
        ],
    )

    replace_range(
        doc,
        224,
        258,
        [
            "1.5 Phương pháp tiếp cận",
            "Đề tài được triển khai theo hướng kết hợp phân tích nghiệp vụ thương mại điện tử, thiết kế kiến trúc full-stack và thực nghiệm trên hệ thống đang chạy.",
            "1. Khảo sát nhu cầu người dùng và seller",
            "Quan sát luồng mua sắm trên các nền tảng thương mại điện tử phổ biến để xác định các chức năng cần có như tìm kiếm, catalog, giỏ hàng, checkout, theo dõi đơn, đánh giá và hoàn tiền.",
            "Phân tích nhu cầu seller trong việc quản lý shop, sản phẩm, variant, tồn kho, vận chuyển, review, coupon, marketing và tài chính.",
            "Phân tích nhu cầu admin trong kiểm duyệt seller/sản phẩm, quản lý đơn hàng, xử lý return, xem analytics, train recommendation và duyệt payout.",
            "2. Lựa chọn giải pháp hệ thống",
            "Hệ thống được thiết kế theo kiến trúc client-server gồm user-fe, admin-FE, BE và PostgreSQL. Frontend sử dụng React + Vite, backend sử dụng FastAPI và Prisma Python.",
            "Dữ liệu hành vi được ghi nhận từ các hành động VIEW, CLICK, ADD_TO_CART và PURCHASE để phục vụ mô hình gợi ý.",
            "3. Lựa chọn giải pháp AI Recommendation",
            "Đề tài sử dụng Item-based Collaborative Filtering làm mô hình phục vụ chính cho recommendation.",
            "Bên cạnh đó, hệ thống có thêm Learning-to-Rank với LightGBM/XGBoost hoặc fallback tuyến tính, cùng Two-Tower/NCF để thử nghiệm mở rộng.",
            "Hệ thống đánh giá recommendation bằng chronological holdout với các chỉ số HitRate@K, NDCG@K; CTR và Conversion được tính từ bảng hành vi.",
            "4. Phương pháp triển khai",
            "Phân tích yêu cầu và mô hình dữ liệu.",
            "Xây dựng backend API theo module nghiệp vụ.",
            "Xây dựng user frontend, seller center và admin frontend.",
            "Tích hợp thanh toán, vận chuyển, hoàn tiền, finance và recommendation.",
            "Seed dữ liệu marketplace có shop, sản phẩm, ảnh, order, review và behavior để kiểm thử.",
            "Kiểm thử chức năng, kiểm thử build frontend và kiểm thử một số luồng backend quan trọng.",
        ],
    )

    replace_range(
        doc,
        164,
        223,
        [
            "1.4 Nội dung thực hiện",
            "Nghiên cứu cơ sở lý thuyết và công nghệ liên quan",
            "Nghiên cứu tổng quan về thương mại điện tử, marketplace đa người bán và các luồng nghiệp vụ mua - bán - vận chuyển - hoàn tiền.",
            "Tìm hiểu hệ thống gợi ý sản phẩm, bao gồm collaborative filtering, content-based filtering, hybrid recommendation và learning-to-rank.",
            "Tìm hiểu các công nghệ sử dụng trong đề tài: React, TypeScript, Vite, Tailwind CSS, FastAPI, Prisma Python, PostgreSQL, Cloudinary, Docker và Redis.",
            "2. Phân tích và đặc tả yêu cầu hệ thống",
            "Xác định các nhóm vai trò chính: Customer, Seller và Admin.",
            "Phân tích yêu cầu mua hàng, bán hàng, quản trị, tài chính và gợi ý sản phẩm.",
            "Xây dựng use case và các luồng nghiệp vụ chính như checkout, thanh toán, vận chuyển theo shop package, review, return/refund, commission và payout.",
            "3. Thiết kế hệ thống",
            "Thiết kế kiến trúc tổng thể gồm user-fe, admin-FE, BE, PostgreSQL và AI module.",
            "Thiết kế cơ sở dữ liệu quan hệ bằng Prisma schema.",
            "Thiết kế các bảng quan trọng như User, Shop, Product, ProductVariant, Order, OrderShopPackage, Payment, Review, Wishlist, UserBehavior, PlatformCommission và SellerPayout.",
            "Thiết kế REST API cho các module nghiệp vụ.",
            "Thiết kế giao diện người dùng, seller center và admin center.",
            "4. Xây dựng hệ thống thương mại điện tử",
            "Triển khai frontend bằng React + Vite: trang chủ, catalog, chi tiết sản phẩm, giỏ hàng, checkout, đơn hàng, wishlist, review, return request và recommendation.",
            "Triển khai Seller Center: dashboard, sản phẩm, kho/variant, đơn hàng, vận chuyển, trả hàng, review, coupon, marketing, khách hàng, phân tích, tài chính và cài đặt shop.",
            "Triển khai Admin Center: dashboard, user, shop/seller application, product moderation, order, return, review, marketing, inventory, finance, analytics, audit và security.",
            "Triển khai backend bằng FastAPI: auth, cart, product, order, shipment, payment, return, review, wishlist, finance, analytics, upload, seller, admin và support.",
            "5. Xây dựng và tích hợp hệ thống AI đề xuất",
            "Thu thập hành vi người dùng từ bảng UserBehavior.",
            "Quy đổi hành vi thành trọng số tương tác cho model.",
            "Huấn luyện Item-based Collaborative Filtering và các mô hình ranking thử nghiệm.",
            "Lưu model bằng joblib và cung cấp API recommendation cho frontend.",
            "Đánh giá bằng HitRate@K, NDCG@K, CTR và Conversion.",
            "6. Kiểm thử và đánh giá hệ thống",
            "Kiểm thử các luồng checkout, order package, finance commission, analytics/recommendation và build frontend.",
            "Kiểm tra các lỗi thực tế như upload ảnh 400, wishlist serialize, xác nhận đã nhận hàng 403 và trạng thái package đa shop.",
            "7. Đóng gói và triển khai hệ thống",
            "Chuẩn bị Docker, nginx, biến môi trường và cấu hình để triển khai production.",
            "8. Hoàn thiện báo cáo và chuẩn bị bảo vệ",
            "Tổng hợp kết quả thực hiện, cập nhật báo cáo theo phiên bản hệ thống hiện tại và chuẩn bị kịch bản demo.",
        ],
    )

    replace_range(
        doc,
        88,
        121,
        [
            "1.2.2 Mục tiêu cụ thể",
            "Để đạt được mục tiêu tổng quát, đề tài hướng đến các mục tiêu cụ thể sau:",
            "Phân tích và xác định yêu cầu hệ thống",
            "Khảo sát nghiệp vụ thương mại điện tử đa người bán, xác định yêu cầu của khách hàng, seller và admin.",
            "Xây dựng đặc tả yêu cầu cho các luồng mua hàng, bán hàng, vận chuyển, trả hàng, tài chính và gợi ý sản phẩm.",
            "Thiết kế kiến trúc và cơ sở dữ liệu hệ thống",
            "Thiết kế kiến trúc tách biệt user frontend, admin frontend, backend API, database và AI module.",
            "Xây dựng mô hình cơ sở dữ liệu PostgreSQL bằng Prisma schema, đảm bảo lưu trữ đầy đủ dữ liệu người dùng, shop, sản phẩm, biến thể, đơn hàng, thanh toán, vận chuyển, review, wishlist, hành vi, hoa hồng và payout.",
            "Xây dựng hệ thống thương mại điện tử đa người bán",
            "Triển khai đăng ký, đăng nhập, quản lý tài khoản, địa chỉ, wishlist và review.",
            "Triển khai quản lý sản phẩm, danh mục, hình ảnh, biến thể, thuộc tính, tồn kho và tìm kiếm/lọc sản phẩm.",
            "Triển khai giỏ hàng, checkout, coupon, flash sale, thanh toán COD/VNPay/MoMo và theo dõi đơn hàng.",
            "Triển khai package theo từng shop trong cùng một hóa đơn để mỗi shop có trạng thái vận chuyển riêng.",
            "Triển khai Seller Center",
            "Xây dựng giao diện seller quản lý shop, sản phẩm, variant, kho, đơn hàng, vận chuyển, return/refund, review, coupon, marketing, khách hàng, phân tích và tài chính.",
            "Triển khai Admin Center",
            "Xây dựng giao diện admin quản lý user, seller application, shop, product moderation, order, return, review, marketing, analytics, finance, audit và security.",
            "Thu thập và xử lý dữ liệu hành vi người dùng",
            "Ghi nhận các hành vi VIEW, CLICK, ADD_TO_CART và PURCHASE.",
            "Quy đổi hành vi và dữ liệu đơn hàng thành trọng số tương tác phục vụ recommendation.",
            "Xây dựng và huấn luyện mô hình AI đề xuất",
            "Áp dụng Item-based Collaborative Filtering làm mô hình gợi ý chính.",
            "Bổ sung Learning-to-Rank bằng LightGBM/XGBoost hoặc fallback tuyến tính, đồng thời chuẩn bị Two-Tower/NCF để thử nghiệm.",
            "Triển khai API recommendation, xử lý cold start bằng sản phẩm phổ biến, sản phẩm cùng danh mục hoặc sản phẩm tương tự.",
            "Đánh giá hiệu quả hệ thống đề xuất",
            "Đo lường HitRate@K, NDCG@K bằng chronological holdout; tính CTR và Conversion từ dữ liệu hành vi.",
            "Hoàn thiện tài chính nền tảng",
            "Tính phí hoa hồng 3-7% theo giá trị đơn/package, lưu PlatformCommission, hiển thị ledger admin và duyệt SellerPayout.",
            "Kiểm thử và đóng gói hệ thống",
            "Kiểm thử các luồng chính, build frontend, train recommendation và chuẩn bị Docker/nginx cho triển khai.",
        ],
    )

    replacements = {
        "Nest.js": "FastAPI",
        "NestJS": "FastAPI",
        "Next.js": "React + Vite",
        "TypeORM": "Prisma ORM",
        "MySQL": "PostgreSQL",
        "SQLAlchemy": "Prisma ORM",
        "Passport.js + passport-google-oauth20": "JWT, refresh token và RBAC",
        "Passport.js": "JWT/refresh token",
        "Google OAuth 2.0": "JWT và refresh token",
        "Google OAuth": "JWT",
        "Google Client ID/Secret": "JWT secret và database connection",
        "SVD": "Item-based Collaborative Filtering và Learning-to-Rank",
        "CollabPMS": "hệ thống thương mại điện tử đa người bán",
        "Kanban": "quy trình xử lý đơn hàng",
        "task": "đơn hàng",
        "Task": "Đơn hàng",
        "project": "shop",
        "Project": "Shop",
        "Trello/Jira": "Shopee Seller Center và các trang quản trị thương mại điện tử",
        "Material-UI (MUI)": "Tailwind CSS và component UI",
        "Framer Motion": "CSS transition/Tailwind",
        "class-validator": "Pydantic schema",
        "TypeSc": "Prisma Python",
        "Swager": "Swagger",
        "user_interactions": "UserBehavior",
        "Precision@K và Recall@K": "HitRate@K và NDCG@K",
    }
    replace_in_all_text(doc, replacements)
    update_table_text(doc)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
