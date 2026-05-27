from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


INPUT = Path(r"D:\DATN\document_work\10_12522084_DaoAnhQuan_BaoCao_DA_DaChinh_Final_ThamKhao.docx")
OUTPUT = Path(r"D:\DATN\document_work\10_12522084_DaoAnhQuan_BaoCao_DA_DaChinh_Final_MoRongNoiDung.docx")


def iter_all_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def set_paragraph_text(paragraph, text):
    style = paragraph.style
    clear_paragraph(paragraph)
    paragraph.style = style
    paragraph.add_run(text)


def paragraph_by_text(doc, needle, style=None):
    for paragraph in doc.paragraphs:
        if needle in paragraph.text and (style is None or paragraph.style.name == style):
            return paragraph
    raise ValueError(f"Paragraph containing {needle!r} with style {style!r} not found")


def add_paragraph_after(cursor, text="", style=None):
    paragraph = doc.add_paragraph()
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    cursor.addnext(paragraph._p)
    return paragraph._p


def set_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_table_after(cursor, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_header(table.rows[0])
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    cursor.addnext(table._tbl)
    return table._tbl


doc = Document(INPUT)

replacements = {
    "Thông qua việc tích hợp thuật toán Collaborative Filtering trong hệ thống đề xuất, đề tài hướng đến giải quyết bài toán quá tải thông tin trong môi trường thương mại điện tử, nơi số lượng sản phẩm ngày càng lớn và người dùng khó lựa chọn sản phẩm phù hợp.":
        "Thông qua việc tích hợp item-based collaborative filtering, learning-to-rank và các chiến lược fallback cho người dùng mới, đề tài hướng đến giải quyết bài toán quá tải thông tin trong môi trường thương mại điện tử, nơi số lượng sản phẩm ngày càng lớn và người dùng khó lựa chọn sản phẩm phù hợp.",
    "Thông qua việc tích hợp thuật toán Collaborative Filtering trong hệ thống đề xuất, đề tài hướng đến giải quyết bài toán quá tải thông tin trong môi trường thương mại điện tử, nơi số lượng sản phẩm lớn khiến người dùng gặp khó khăn trong việc lựa chọn. Việc áp dụng AI không chỉ giúp tối ưu quá trình tìm kiếm sản phẩm mà còn tạo nền tảng cho các nghiên cứu và phát triển sâu hơn về hệ thống gợi ý trong tương lai.":
        "Thông qua việc tích hợp item-based collaborative filtering, learning-to-rank và cơ chế fallback, đề tài hướng đến giải quyết bài toán quá tải thông tin trong môi trường thương mại điện tử, nơi số lượng sản phẩm lớn khiến người dùng gặp khó khăn trong việc lựa chọn. Việc áp dụng AI không chỉ giúp tối ưu quá trình tìm kiếm sản phẩm mà còn tạo nền tảng cho các nghiên cứu và phát triển sâu hơn về hệ thống gợi ý trong tương lai.",
    "Các mô hình và thuật toán hệ thống gợi ý (Recommendation System), đặc biệt là phương pháp Collaborative Filtering.":
        "Các mô hình và thuật toán hệ thống gợi ý (Recommendation System), gồm collaborative filtering, content-based filtering, hybrid recommendation và learning-to-rank.",
    "Thuật toán phân rã ma trận Item-based Collaborative Filtering và Learning-to-Rank (Singular Value Decomposition) trong việc xây dựng hệ thống đề xuất.":
        "Phương pháp item-based collaborative filtering, learning-to-rank và cơ chế fallback theo sản phẩm phổ biến, cùng danh mục hoặc sản phẩm tương tự trong việc xây dựng hệ thống đề xuất.",
    "Frontend được triển khai bằng React":
        "Frontend được tách thành user-fe và admin-FE, triển khai bằng React, TypeScript và Vite",
    "Backend được triển khai trên các nền tảng như Render hoặc DigitalOcean":
        "Backend được triển khai bằng FastAPI; trong phạm vi đồ án chạy trên môi trường cục bộ và chuẩn bị cấu hình Docker/nginx để triển khai production",
    "Cơ sở dữ liệu sử dụng PostgreSQL đặt trên server cloud":
        "Cơ sở dữ liệu sử dụng PostgreSQL, có thể chạy cục bộ trong giai đoạn kiểm thử hoặc tách riêng trên server/database managed khi triển khai production",
    "Trong tương lai, nền tảng có thể tiếp tục được mở rộng với các chức năng nâng cao như tích hợp thanh toán trực tuyến, hệ thống trò chuyện giữa người mua và người bán, hoặc phân tích dữ liệu để hỗ trợ quyết định kinh doanh.":
        "Trong tương lai, nền tảng có thể tiếp tục được mở rộng theo hướng hoàn thiện thanh toán production, tự động hóa chuyển khoản payout, hệ thống trò chuyện giữa người mua và người bán, hoặc phân tích dữ liệu để hỗ trợ quyết định kinh doanh.",
    "Xác thực: JWT + refresh token và refresh token":
        "Xác thực: JWT access token + refresh token",
    "Kiến trúc tổng thể của hệ thống hệ thống thương mại điện tử đa người bán":
        "Kiến trúc tổng thể của hệ thống thương mại điện tử đa người bán",
    "Hệ thống hỗ trợ đăng nhập nhanh thông qua JWT":
        "Hệ thống hỗ trợ đăng nhập bằng email/mật khẩu, JWT access token và refresh token",
}
for paragraph in iter_all_paragraphs(doc):
    original = paragraph.text
    updated = original
    for old, new in replacements.items():
        updated = updated.replace(old, new)
    if updated != original:
        set_paragraph_text(paragraph, updated)


# 1. Add a clear novelty section after the reason for choosing the topic.
reason_end = paragraph_by_text(doc, "được lựa chọn nhằm nghiên cứu, thiết kế và triển khai")
cursor = reason_end._p
intro_items = [
    ("Điểm mới và đóng góp chính của đề tài", "MucNho"),
    (
        "Điểm khác biệt của đề tài so với một website bán hàng thông thường là hệ thống được thiết kế theo mô hình marketplace đa người bán. Một đơn hàng của khách hàng có thể chứa sản phẩm từ nhiều shop, nhưng quá trình xử lý, vận chuyển, hoàn tiền và đối soát tài chính phải được tách riêng theo từng shop.",
        "Normal1",
    ),
    (
        "Hệ thống không chỉ dừng ở chức năng mua hàng cơ bản mà còn có Seller Center cho người bán, Admin Center cho quản trị viên, module tài chính tính phí hoa hồng 3-7%, yêu cầu rút tiền seller và dashboard recommendation để train/evaluate mô hình AI.",
        "Normal1",
    ),
    (
        "Về mặt AI, hệ thống ghi nhận hành vi VIEW, CLICK, ADD_TO_CART và PURCHASE, sau đó sử dụng dữ liệu này để gợi ý sản phẩm. Khi dữ liệu ít hoặc người dùng mới chưa có lịch sử, hệ thống có cơ chế fallback bằng sản phẩm phổ biến, sản phẩm cùng danh mục hoặc sản phẩm liên quan để vẫn trả được kết quả gợi ý.",
        "Normal1",
    ),
    (
        "Về mặt nghiệp vụ vận hành, đề tài bổ sung các luồng dễ phát sinh lỗi trong thực tế như upload ảnh sản phẩm/variant, review sau mua, return/refund có bằng chứng media, tracking trạng thái theo shop package, kiểm soát lỗi 403 theo quyền truy cập và đối soát payout cho seller.",
        "Normal1",
    ),
]
for text, style in intro_items:
    cursor = add_paragraph_after(cursor, text, style)


# 2. Add role analysis and business rules in the requirements chapter.
business_issues = paragraph_by_text(doc, "Hệ thống cần ghi nhận hành vi người dùng")
cursor = business_issues._p
role_caption = add_paragraph_after(cursor, "Bảng 3.x: Phân tích vai trò và phạm vi thao tác chính", "Caption")
role_table = add_table_after(
    role_caption,
    ["Vai trò", "Mục tiêu sử dụng", "Phạm vi thao tác chính", "Ràng buộc quyền"],
    [
        (
            "Customer",
            "Mua hàng, theo dõi đơn và nhận gợi ý sản phẩm.",
            "Đăng ký/đăng nhập, quản lý địa chỉ, xem sản phẩm, giỏ hàng, checkout, thanh toán, wishlist, review, return/refund.",
            "Chỉ được xem đơn hàng của mình, chỉ xác nhận nhận hàng hoặc review với đơn/package thuộc tài khoản của mình.",
        ),
        (
            "Seller",
            "Vận hành shop và xử lý phần đơn hàng thuộc shop.",
            "Quản lý shop, sản phẩm, variant, tồn kho, đơn hàng, vận chuyển, review, coupon, marketing, tài chính và payout.",
            "Chỉ được thao tác trên dữ liệu shop của mình; không được xem/sửa dữ liệu seller khác.",
        ),
        (
            "Admin",
            "Quản trị toàn bộ nền tảng và đảm bảo vận hành đúng quy định.",
            "Quản lý user/shop/product/order/return, kiểm duyệt, analytics, train/evaluate AI, cấu hình hoa hồng, xem ledger và duyệt payout.",
            "Có quyền quản trị tập trung nhưng các thao tác nhạy cảm cần được ghi log để phục vụ truy vết.",
        ),
    ],
)

functional_heading = paragraph_by_text(doc, "Các yêu cầu chức năng", style="Tiểu mục")
previous = functional_heading._p.getprevious()
cursor = previous if previous is not None else functional_heading._p
for text, style in [
    ("Quy tắc nghiệp vụ quan trọng", "Tiểu mục"),
    (
        "Mỗi order có thể được chia thành nhiều OrderShopPackage theo shop. Trạng thái tổng của order được đồng bộ từ trạng thái các package, nhưng seller chỉ cập nhật package của shop mình.",
        "Nội dung",
    ),
    (
        "Khách hàng chỉ được xác nhận đã nhận hàng với package thuộc đơn hàng của chính họ. Nếu token hợp lệ nhưng tài nguyên không thuộc quyền truy cập, backend phải trả lỗi 403 Forbidden thay vì cho cập nhật sai dữ liệu.",
        "Nội dung",
    ),
    (
        "Phí nền tảng chỉ được ghi nhận khi order item/package đạt trạng thái doanh thu hợp lệ như DELIVERED hoặc COMPLETED. Các trường hợp hủy, trả hàng hoặc hoàn tiền phải cập nhật commission sang trạng thái CANCELLED/REFUNDED tương ứng.",
        "Nội dung",
    ),
    (
        "Seller chỉ được tạo yêu cầu rút tiền nếu số tiền nhỏ hơn hoặc bằng availableBalance. Các payout ở trạng thái PENDING/PROCESSING phải được trừ khỏi số dư khả dụng để tránh rút vượt.",
        "Nội dung",
    ),
    (
        "Review sản phẩm chỉ nên được tạo sau khi người dùng đã mua hàng; media upload phải dùng endpoint upload hợp lệ và lưu URL sau khi upload thành công.",
        "Nội dung",
    ),
]:
    cursor = add_paragraph_after(cursor, text, style)


# 3. Expand finance design with explicit commission tiers and status handling.
finance_heading = paragraph_by_text(doc, "Thiết kế tài chính, hoa hồng và payout", style="Tiểu mục")
cursor = finance_heading._p
while cursor.getnext() is not None:
    cursor = cursor.getnext()
    # Stop after the existing finance section before architecture.
    if cursor.text and cursor.text.strip().startswith("Thiết kế kiến trúc"):
        cursor = cursor.getprevious()
        break

for text, style in [
    ("Bảng 3.x: Mức hoa hồng mặc định theo giá trị package/shop", "Caption"),
]:
    cursor = add_paragraph_after(cursor, text, style)
commission_table = add_table_after(
    cursor,
    ["Giá trị hàng hóa theo shop/package", "Tỉ lệ hoa hồng mặc định", "Ý nghĩa nghiệp vụ"],
    [
        ("Từ 5.000.000đ trở lên", "3%", "Đơn giá trị cao được giảm phí để khuyến khích seller bán sản phẩm giá trị lớn."),
        ("Từ 2.000.000đ đến dưới 5.000.000đ", "4%", "Mức phí trung gian cho các đơn có giá trị khá cao."),
        ("Từ 1.000.000đ đến dưới 2.000.000đ", "5%", "Mức phí phổ biến cho nhóm đơn trung bình."),
        ("Từ 500.000đ đến dưới 1.000.000đ", "6%", "Mức phí cho nhóm đơn nhỏ hơn."),
        ("Dưới 500.000đ", "7%", "Mức phí cao nhất do chi phí vận hành tương đối lớn so với giá trị đơn."),
    ],
)
cursor = commission_table
for text, style in [
    (
        "Ngoài mức mặc định theo giá trị đơn, hệ thống còn hỗ trợ cấu hình hoa hồng riêng theo shop hoặc theo danh mục. Dù cấu hình theo cách nào, backend vẫn chặn tỉ lệ ngoài khoảng 3% đến 7% để tránh nhập sai dữ liệu tài chính.",
        "Nội dung",
    ),
    (
        "Các trạng thái tài chính được tách rõ: PENDING cho khoản phí chưa đủ điều kiện chốt, EARNED cho khoản phí đã phát sinh sau khi package hoàn tất, SETTLED cho khoản đã đối soát, REFUNDED/CANCELLED cho trường hợp hoàn trả hoặc hủy. Cách tách trạng thái này giúp admin xem đúng doanh thu nền tảng và seller biết số dư nào có thể rút.",
        "Nội dung",
    ),
]:
    cursor = add_paragraph_after(cursor, text, style)


# 4. Add detailed design sections before database design.
architecture_summary = paragraph_by_text(doc, "Sơ đồ kiến trúc tổng thể của hệ thống")
cursor = architecture_summary._p
design_sections = [
    ("Thiết kế luồng đơn hàng nhiều shop", "Tiểu mục"),
    (
        "Khi khách hàng checkout, backend đọc các CartItem trong giỏ hàng và nhóm theo shopId. Từ đó hệ thống tạo một Order tổng để giữ thông tin người mua, địa chỉ giao hàng, thanh toán và tổng tiền.",
        "Nội dung",
    ),
    (
        "Với mỗi shop xuất hiện trong đơn, hệ thống tạo một OrderShopPackage. Package này lưu shopId, trạng thái xử lý, mã vận đơn, thời điểm giao hàng và thời điểm khách xác nhận đã nhận hàng. Nhờ đó, cùng một hóa đơn có thể có nhiều trạng thái vận chuyển khác nhau.",
        "Nội dung",
    ),
    (
        "Seller chỉ được cập nhật trạng thái package của shop mình theo các bước hợp lệ như PENDING, CONFIRMED, PROCESSING, READY_TO_SHIP, SHIPPED, IN_TRANSIT, OUT_FOR_DELIVERY, DELIVERED. Customer xác nhận nhận hàng sẽ chuyển package sang COMPLETED và là điều kiện để ghi nhận số dư seller.",
        "Nội dung",
    ),
    (
        "Trạng thái Order tổng được đồng bộ từ trạng thái các package. Nếu các package còn khác trạng thái, hệ thống chọn trạng thái đại diện theo thứ tự ưu tiên nghiệp vụ để frontend vẫn hiển thị được tiến độ đơn hàng.",
        "Nội dung",
    ),
    ("Thiết kế pipeline gợi ý sản phẩm", "Tiểu mục"),
    (
        "Pipeline recommendation bắt đầu từ việc ghi nhận UserBehavior khi người dùng xem, click, thêm vào giỏ hoặc mua sản phẩm. Mỗi hành vi được lưu cùng userId, productId, sessionId, thời lượng và metadata để có thể phân tích về sau.",
        "Nội dung",
    ),
    (
        "FeatureBuilder quy đổi các hành vi thành interaction có trọng số. PURCHASE có giá trị cao hơn ADD_TO_CART, CLICK và VIEW vì phản ánh ý định mua rõ hơn. Các interaction này là đầu vào cho mô hình item-based collaborative filtering và các mô hình ranking.",
        "Nội dung",
    ),
    (
        "Khi gọi API recommendation, hệ thống ưu tiên gợi ý cá nhân hóa theo lịch sử của user. Nếu user chưa đủ dữ liệu, hệ thống fallback sang sản phẩm phổ biến, sản phẩm cùng danh mục, sản phẩm liên quan đến context product hoặc kết quả tìm kiếm.",
        "Nội dung",
    ),
    (
        "Dashboard admin có chức năng train và evaluate mô hình. Các chỉ số HitRate@K và NDCG@K được tính bằng chronological holdout, còn CTR và Conversion được tính từ số view/click/purchase thực tế. Khi dữ liệu ít, báo cáo cần giải thích rằng metric có thể bằng 0 và chưa phản ánh chất lượng production.",
        "Nội dung",
    ),
    ("Thiết kế bảo mật và phân quyền", "Tiểu mục"),
    (
        "Hệ thống sử dụng JWT access token cho request API và refresh token để làm mới phiên đăng nhập. Storefront và admin có luồng login/refresh/logout riêng để hạn chế dùng nhầm phiên giữa người dùng thường và quản trị viên.",
        "Nội dung",
    ),
    (
        "RBAC được áp dụng tại backend thay vì chỉ ẩn nút trên frontend. Customer, Seller và Admin được kiểm tra ở từng endpoint quan trọng; seller phải được đối chiếu shop owner, còn customer phải được đối chiếu userId của order.",
        "Nội dung",
    ),
    (
        "Các thao tác nhạy cảm như cập nhật trạng thái đơn, duyệt return, duyệt payout, kiểm duyệt sản phẩm và đăng nhập admin cần ghi audit/security log để admin có thể truy vết khi phát sinh lỗi hoặc khiếu nại.",
        "Nội dung",
    ),
    ("Thiết kế upload media", "Tiểu mục"),
    (
        "Media trong hệ thống gồm ảnh sản phẩm, ảnh variant, ảnh review và bằng chứng trả hàng/hoàn tiền. Backend tiếp nhận file upload, đẩy lên Cloudinary và chỉ lưu URL cùng metadata cần thiết vào PostgreSQL.",
        "Nội dung",
    ),
    (
        "Cách thiết kế này giúp database không phải lưu trực tiếp file nhị phân lớn, frontend tải ảnh qua CDN nhanh hơn và các module sản phẩm/review/return có thể dùng chung một cơ chế upload.",
        "Nội dung",
    ),
]
for text, style in design_sections:
    cursor = add_paragraph_after(cursor, text, style)


# 5. Expand implemented feature sections for seller/admin and business workflows.
system_functions = paragraph_by_text(doc, "Các chức năng hệ thống", style="Tiểu mục")
cursor = system_functions._p
seller_admin_sections = [
    ("Các chức năng nghiệp vụ phân hệ Seller Center", "Tiểu mục"),
    (
        "Seller Center được xây dựng để người bán tự vận hành shop mà không cần thao tác trực tiếp trong trang admin. Header seller hiển thị tên shop, hỗ trợ chuyển về trang người dùng thường và truy cập trang xem shop để kiểm tra cách khách hàng nhìn thấy gian hàng.",
        "Nội dung",
    ),
    (
        "Seller có thể quản lý sản phẩm, ảnh sản phẩm, biến thể, tồn kho và giá bán. Với sản phẩm có nhiều variant, mỗi variant có giá, tồn kho và ảnh riêng để tránh lỗi mất ảnh khi hiển thị trang chi tiết sản phẩm.",
        "Nội dung",
    ),
    (
        "Seller quản lý đơn hàng theo package của shop mình. Trong cùng một hóa đơn, mỗi seller chỉ nhìn thấy phần hàng thuộc shop của mình và cập nhật trạng thái vận chuyển riêng, giúp tránh nhầm lẫn khi nhiều shop cùng tham gia một đơn hàng.",
        "Nội dung",
    ),
    (
        "Trang tài chính seller hiển thị doanh thu hoàn tất, doanh thu đang chờ, phí nền tảng, hoàn/hủy, số dư có thể rút và lịch sử payout. Seller chỉ gửi payout khi số dư khả dụng đủ và có thể theo dõi trạng thái PENDING, PROCESSING, PAID, FAILED hoặc CANCELLED.",
        "Nội dung",
    ),
    ("Các chức năng nghiệp vụ phân hệ Admin Center", "Tiểu mục"),
    (
        "Admin Center tập trung vào quản trị nền tảng: quản lý user, shop, sản phẩm, kiểm duyệt nội dung, đơn hàng, trả hàng/hoàn tiền, audit/security log và báo cáo vận hành.",
        "Nội dung",
    ),
    (
        "Với tài chính, admin xem được tổng doanh thu gross, phí hoa hồng nền tảng, số dư seller, payout đang chờ và payout đã thanh toán. Admin có thể cấu hình hoa hồng theo shop hoặc danh mục trong khoảng 3-7% và duyệt/từ chối yêu cầu rút tiền.",
        "Nội dung",
    ),
    (
        "Với AI recommendation, admin có màn hình train/evaluate, theo dõi HitRate@K, NDCG@K, CTR, Conversion, số lượng interaction và số user được đánh giá. Màn hình này giúp chứng minh hệ thống không chỉ hardcode gợi ý mà có pipeline dữ liệu và đánh giá mô hình.",
        "Nội dung",
    ),
    ("Các luồng nghiệp vụ đã hoàn thiện sau khi chỉnh sửa", "Tiểu mục"),
    (
        "Luồng upload ảnh được chuẩn hóa để sản phẩm và variant luôn lưu đúng URL ảnh. Điều này giải quyết tình trạng trang chi tiết variant bị mất ảnh khi dữ liệu variant image không đồng bộ với product image.",
        "Nội dung",
    ),
    (
        "Luồng xác nhận đã nhận hàng được kiểm tra lại theo quyền Customer và theo package. Nếu customer nhấn xác nhận package không thuộc đơn của mình hoặc sai trạng thái, backend trả lỗi phù hợp thay vì cập nhật nhầm dữ liệu.",
        "Nội dung",
    ),
    (
        "Luồng review được cải thiện về giao diện và nghiệp vụ: review gắn với sản phẩm/order đã mua, có điểm đánh giá, nội dung nhận xét và có thể kèm media, giúp dữ liệu phản hồi hữu ích hơn cho khách hàng khác.",
        "Nội dung",
    ),
]
previous = cursor.getprevious()
insert_cursor = previous if previous is not None else cursor
for text, style in seller_admin_sections:
    insert_cursor = add_paragraph_after(insert_cursor, text, style)


# 6. Add testing content for the newly emphasized workflows.
result_heading = paragraph_by_text(doc, "Kết quả kiểm thử")
previous = result_heading._p.getprevious()
cursor = previous if previous is not None else result_heading._p
testing_items = [
    ("Kiểm thử nghiệp vụ đặc thù của marketplace", "Nội dung"),
    (
        "Kiểm thử order nhiều shop: tạo giỏ hàng có sản phẩm từ nhiều shop, checkout một lần, kiểm tra số lượng OrderShopPackage được tạo, kiểm tra seller A không xem/cập nhật được package của seller B và trạng thái order tổng đồng bộ đúng.",
        "Nội dung",
    ),
    (
        "Kiểm thử tài chính: cập nhật package sang DELIVERED/COMPLETED, chạy đồng bộ commission, kiểm tra PlatformCommission sinh đúng theo tỉ lệ 3-7%, kiểm tra availableBalance đã trừ commission và payout pending.",
        "Nội dung",
    ),
    (
        "Kiểm thử payout: seller tạo yêu cầu rút tiền hợp lệ, admin chuyển trạng thái PROCESSING/PAID hoặc FAILED/CANCELLED, hệ thống không cho đổi payout đã ở trạng thái cuối và không cho rút vượt số dư.",
        "Nội dung",
    ),
    (
        "Kiểm thử recommendation: seed hành vi VIEW/CLICK/ADD_TO_CART/PURCHASE, train model, gọi API recommend và evaluate, kiểm tra hệ thống trả fallback khi user chưa đủ lịch sử và hiển thị metric đúng với dữ liệu hiện có.",
        "Nội dung",
    ),
    (
        "Kiểm thử upload/review/return: upload ảnh sản phẩm và ảnh variant, tạo review sau mua, tạo return request có bằng chứng media và kiểm tra frontend hiển thị đúng ảnh ở các màn hình liên quan.",
        "Nội dung",
    ),
]
for text, style in testing_items:
    cursor = add_paragraph_after(cursor, text, style)


doc.save(OUTPUT)
print(OUTPUT)
